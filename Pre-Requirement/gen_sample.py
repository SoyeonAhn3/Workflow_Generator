# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# 페이지 설정
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# 기본 폰트
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_width(cell, width_twips):
    tcPr = cell._element.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcPr.append(tcW)


def add_text(cell, text, bold=False, size=10, color=None, align=None):
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = color


def set_borders(table, color='D1D5DB'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(borders)


# ================================================================
# 샘플 데이터
# ================================================================
steps = [
    {
        'no': 1,
        'title': '모든 부서 데이터 입력 완료 확인',
        'dept': '전 부서',
        'pt': '60min',
        'tcode': 'FB50',
        'logic': '월 마감 전 모든 부서에서 데이터 입력이 완료되었는지 확인합니다.\n미입력 부서가 있을 경우 담당자에게 연락하여 입력을 독촉합니다.',
        'warning': '',
        'note': '마감 기한: 매월 말일 17:00까지',
    },
    {
        'no': 2,
        'title': 'SAP A데이터 확인',
        'dept': '재무팀',
        'pt': '30min',
        'tcode': 'KSB5',
        'logic': 'SAP 메인 화면 접속 후 A데이터 메뉴로 이동합니다.\n반드시 당월 기준으로 조회해야 합니다.',
        'warning': '반드시 당월 기준으로 조회할 것',
        'note': '',
    },
    {
        'no': 3,
        'title': 'B Tcode 전표 생성',
        'dept': '재무팀',
        'pt': '20min',
        'tcode': 'F-02',
        'logic': 'Tcode B 입력 후 진입하여 화면 하단 [전표 생성] 버튼을 클릭합니다.\n생성 완료 메시지를 반드시 확인합니다.',
        'warning': '생성 완료 메시지 반드시 확인',
        'note': '전표 번호를 별도로 기록해 둘 것',
    },
]

# ================================================================
# 제목
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('단계별 상세 설명 (샘플 템플릿)')
run.font.name = '맑은 고딕'
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p2 = doc.add_paragraph()
run2 = p2.add_run('ProcessFlow 템플릿 참고 · 프로세스 설명란 디자인 샘플')
run2.font.name = '맑은 고딕'
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

# ================================================================
# 단계별 출력
# ================================================================
for step in steps:
    # ── 단계 타이틀 배너 (인디고 번호 + 연보라 제목) ──
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)

    c0 = t.cell(0, 0)
    set_cell_shading(c0, '4F46E5')
    set_cell_width(c0, 600)
    add_text(c0, str(step['no']), bold=True, size=12,
             color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

    c1 = t.cell(0, 1)
    set_cell_shading(c1, 'EEF2FF')
    set_cell_width(c1, 8400)
    add_text(c1, step['title'], bold=True, size=12,
             color=RGBColor(0x4F, 0x46, 0xE5))

    # ── 메타 행 (담당부서 / 처리시간 / 화면명) ──
    t2 = doc.add_table(rows=1, cols=6)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t2)

    meta = [
        ('담당 부서', step['dept'], 'F8FAFC', 'FFFFFF'),
        ('처리시간', step['pt'], 'F8FAFC', 'FFFFFF'),
        ('화면명 (T-Code)', step['tcode'], 'F8FAFC', 'E0F2FE'),
    ]
    for i, (label, val, lbg, vbg) in enumerate(meta):
        lc = t2.cell(0, i * 2)
        set_cell_shading(lc, lbg)
        set_cell_width(lc, 1200)
        add_text(lc, label, bold=True, size=9,
                 color=RGBColor(0x64, 0x64, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

        vc = t2.cell(0, i * 2 + 1)
        set_cell_shading(vc, vbg)
        set_cell_width(vc, 1800)
        add_text(vc, val, size=10, color=RGBColor(0x40, 0x40, 0x40))

    # ── 업무 설명 ──
    t3 = doc.add_table(rows=1, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t3)

    lc3 = t3.cell(0, 0)
    set_cell_shading(lc3, 'F8FAFC')
    set_cell_width(lc3, 1200)
    add_text(lc3, '업무 설명', bold=True, size=9,
             color=RGBColor(0x64, 0x64, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

    vc3 = t3.cell(0, 1)
    set_cell_width(vc3, 7800)
    lines = step['logic'].split('\n')
    first_run = vc3.paragraphs[0].add_run(lines[0])
    first_run.font.name = '맑은 고딕'
    first_run.font.size = Pt(10)
    first_run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    for line in lines[1:]:
        lp = vc3.add_paragraph()
        lr = lp.add_run(line)
        lr.font.name = '맑은 고딕'
        lr.font.size = Pt(10)
        lr.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ── 참고 이미지 ──
    t4 = doc.add_table(rows=1, cols=2)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t4)

    lc4 = t4.cell(0, 0)
    set_cell_shading(lc4, 'F8FAFC')
    set_cell_width(lc4, 1200)
    add_text(lc4, '참고 이미지', bold=True, size=9,
             color=RGBColor(0x64, 0x64, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

    vc4 = t4.cell(0, 1)
    set_cell_width(vc4, 7800)
    add_text(vc4, '(이미지가 여기에 삽입됩니다)', size=9,
             color=RGBColor(0xBF, 0xBF, 0xBF))

    # ── 참고사항 (초록 배경) ──
    if step['note']:
        t5 = doc.add_table(rows=1, cols=2)
        t5.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_borders(t5)

        lc5 = t5.cell(0, 0)
        set_cell_shading(lc5, 'F0FDF4')
        set_cell_width(lc5, 1200)
        add_text(lc5, '참고사항', bold=True, size=9,
                 color=RGBColor(0x16, 0x65, 0x34), align=WD_ALIGN_PARAGRAPH.CENTER)

        vc5 = t5.cell(0, 1)
        set_cell_shading(vc5, 'F0FDF4')
        set_cell_width(vc5, 7800)
        add_text(vc5, step['note'], size=10, color=RGBColor(0x16, 0x65, 0x34))

    # ── 주의사항 (노란 배경) ──
    if step['warning']:
        t6 = doc.add_table(rows=1, cols=2)
        t6.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_borders(t6)

        lc6 = t6.cell(0, 0)
        set_cell_shading(lc6, 'FFFBEB')
        set_cell_width(lc6, 1200)
        add_text(lc6, '주의사항', bold=True, size=9,
                 color=RGBColor(0xD9, 0x77, 0x06), align=WD_ALIGN_PARAGRAPH.CENTER)

        vc6 = t6.cell(0, 1)
        set_cell_shading(vc6, 'FFFBEB')
        set_cell_width(vc6, 7800)
        add_text(vc6, step['warning'], size=10, color=RGBColor(0xD9, 0x77, 0x06))

    # 단계 간 여백
    doc.add_paragraph()

output = 'sample_step_detail_template.docx'
doc.save(output)
print(f'완료: {output}')
